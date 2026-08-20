import os
from typing import Any, Dict, List, Optional

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from . import (
    PARSER_VERSION,
    ParseResult,
    build_text_chunks,
    calculate_file_hash,
    ensure_derived_dir,
    guess_extension_from_content_type,
    load_cached_parse_result,
    summarize_chunks,
    to_rel_upload_path,
    write_cached_parse_result,
)
from .ocr import run_remote_ocr


def parse_docx(
    file_path: str,
    *,
    upload_root: Optional[str] = None,
    owner_id: Any = "system",
    file_hash: Optional[str] = None,
    api_key: Optional[str] = None,
    api_base: Optional[str] = None,
    parse_mode: str = "docx",
) -> ParseResult:
    resolved_hash = file_hash or calculate_file_hash(file_path)
    cached = load_cached_parse_result(upload_root, resolved_hash, parse_mode)
    if cached:
        return cached

    document = DocxDocument(file_path)
    derived_dir = ensure_derived_dir(upload_root, owner_id, resolved_hash, namespace="docx") if upload_root else None
    blocks: List[Dict[str, Any]] = []
    text_parts: List[str] = []
    has_ocr = False
    image_state = {"counter": 0}

    for block in _iter_document_blocks(document):
        if isinstance(block, Paragraph):
            paragraph_text = _normalize_text(block.text)
            heading_level = _detect_heading_level(block)
            if paragraph_text:
                blocks.append(
                    {
                        "block_id": f"body_p_{len(blocks) + 1}",
                        "type": "heading" if heading_level else "paragraph",
                        "text": paragraph_text,
                        "level": heading_level,
                        "source": "body",
                    }
                )
                text_parts.append(paragraph_text)

            image_blocks = _extract_paragraph_images(
                block,
                document,
                derived_dir=derived_dir,
                upload_root=upload_root,
                api_key=api_key,
                api_base=api_base,
                source="body",
                image_state=image_state,
            )
            for image_block in image_blocks:
                if image_block.get("ocr_text"):
                    has_ocr = True
                    text_parts.append(str(image_block.get("ocr_text")))
                blocks.append(image_block)
        elif isinstance(block, Table):
            rows = _extract_table_rows(block)
            if rows:
                table_text = _table_to_text(rows)
                blocks.append(
                    {
                        "block_id": f"body_tbl_{len(blocks) + 1}",
                        "type": "table",
                        "rows": rows,
                        "text": table_text,
                        "source": "body",
                    }
                )
                text_parts.append(table_text)

    _append_header_footer_blocks(document, blocks=blocks, text_parts=text_parts)

    raw_text = "\n\n".join(part for part in text_parts if str(part or "").strip()).strip()
    chunks = _build_docx_chunks(blocks)
    result: ParseResult = {
        "raw_text": raw_text,
        "summary": summarize_chunks(chunks, empty_text="Word 中未提取到可用文本。"),
        "chunks": chunks,
        "structure": {"blocks": blocks},
        "assets": {"block_count": len(blocks)},
        "meta": {
            "parser_version": PARSER_VERSION,
            "file_type": "docx",
            "source_path": file_path,
            "has_ocr": has_ocr,
            "scan_detected": False,
        },
    }
    write_cached_parse_result(upload_root, resolved_hash, parse_mode, result)
    return result


def extract_lines_from_parse_result(result: ParseResult) -> List[str]:
    lines: List[str] = []
    structure_blocks = result.get("structure", {}).get("blocks") if isinstance(result.get("structure"), dict) else []
    if isinstance(structure_blocks, list):
        for block in structure_blocks:
            if not isinstance(block, dict):
                continue
            if str(block.get("type") or "") == "table":
                for row in block.get("rows") or []:
                    row_text = " ".join(str(cell or "").strip() for cell in row if str(cell or "").strip()).strip()
                    if row_text:
                        lines.append(row_text)
                continue
            for key in ("text", "ocr_text"):
                text = _normalize_text(block.get(key))
                if text:
                    lines.extend(line for line in text.splitlines() if _normalize_text(line))
    return [_normalize_text(line) for line in lines if _normalize_text(line)]


def _iter_document_blocks(document: DocxDocument):
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def _normalize_text(value: Any) -> str:
    return " ".join(str(value or "").replace("\xa0", " ").split()).strip()


def _detect_heading_level(paragraph: Paragraph) -> int:
    try:
        style_name = str(paragraph.style.name or "")
    except Exception:
        return 0
    lowered = style_name.lower()
    if not lowered.startswith("heading"):
        return 0
    digits = "".join(ch for ch in style_name if ch.isdigit())
    return int(digits) if digits.isdigit() else 1


def _extract_table_rows(table: Table) -> List[List[str]]:
    rows: List[List[str]] = []
    for row in table.rows:
        row_cells = [_normalize_text(cell.text) for cell in row.cells]
        if any(row_cells):
            rows.append(row_cells)
    return rows


def _table_to_text(rows: List[List[str]]) -> str:
    if not rows:
        return ""
    return "\n".join(" | ".join(cell for cell in row if cell) for row in rows if any(row))


def _extract_paragraph_images(
    paragraph: Paragraph,
    document: DocxDocument,
    *,
    derived_dir: Optional[str],
    upload_root: Optional[str],
    api_key: Optional[str],
    api_base: Optional[str],
    source: str,
    image_state: Dict[str, int],
) -> List[Dict[str, Any]]:
    if derived_dir is None or upload_root is None:
        return []

    image_blocks: List[Dict[str, Any]] = []
    blips = paragraph._element.xpath('.//*[local-name()="blip"]')
    for blip in blips:
        rel_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if not rel_id:
            continue
        image_part = document.part.related_parts.get(rel_id)
        if image_part is None:
            continue
        image_state["counter"] = int(image_state.get("counter", 0)) + 1
        image_index = image_state["counter"]
        extension = guess_extension_from_content_type(getattr(image_part, "content_type", ""), fallback=".png")
        image_abs_path = os.path.join(derived_dir, f"{source}_image_{image_index:03d}{extension}")
        with open(image_abs_path, "wb") as file_obj:
            file_obj.write(image_part.blob)

        ocr_text = ""
        if api_key and api_base:
            try:
                ocr_result = run_remote_ocr(image_abs_path, api_key, api_base)
                ocr_text = _normalize_text(ocr_result.get("raw_text"))
            except Exception:
                ocr_text = ""

        image_blocks.append(
            {
                "block_id": f"{source}_img_{image_index}",
                "type": "image",
                "image_path": to_rel_upload_path(image_abs_path, upload_root),
                "ocr_text": ocr_text,
                "source": source,
            }
        )
    return image_blocks


def _append_header_footer_blocks(document: DocxDocument, *, blocks: List[Dict[str, Any]], text_parts: List[str]) -> None:
    for section in document.sections:
        for source_name, container in (("header", section.header), ("footer", section.footer)):
            seen_texts: set[str] = set()
            for paragraph in container.paragraphs:
                paragraph_text = _normalize_text(paragraph.text)
                if paragraph_text and paragraph_text not in seen_texts:
                    seen_texts.add(paragraph_text)
                    blocks.append(
                        {
                            "block_id": f"{source_name}_{len(blocks) + 1}",
                            "type": source_name,
                            "text": paragraph_text,
                            "source": source_name,
                        }
                    )
                    text_parts.append(paragraph_text)
            for table in container.tables:
                rows = _extract_table_rows(table)
                if not rows:
                    continue
                table_text = _table_to_text(rows)
                blocks.append(
                    {
                        "block_id": f"{source_name}_tbl_{len(blocks) + 1}",
                        "type": "table",
                        "rows": rows,
                        "text": table_text,
                        "source": source_name,
                    }
                )
                text_parts.append(table_text)


def _build_docx_chunks(blocks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    chunks: List[Dict[str, Any]] = []
    for block in blocks:
        block_type = str(block.get("type") or "paragraph")
        if block_type == "table":
            block_text = _table_to_text(block.get("rows") or [])
        else:
            block_text = _normalize_text(block.get("text"))
            if block_type == "image" and not block_text:
                block_text = _normalize_text(block.get("ocr_text"))
        if not block_text:
            continue
        block_chunks = build_text_chunks(block_text, kind=f"docx_{block_type}", extra={"block_id": str(block.get("block_id") or "")})
        for chunk in block_chunks:
            chunk["index"] = len(chunks) + 1
            chunks.append(chunk)
    return chunks
