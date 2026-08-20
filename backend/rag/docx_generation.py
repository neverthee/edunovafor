import hashlib
import json
import os
import re
import time
import uuid
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from werkzeug.utils import secure_filename

from backend.extensions import db
from backend.models.material import Material


DOCX_EXPORT_FILENAME_PREFIX = "lesson-plan"
WORD_TEMPLATE_ROOT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "templates", "lesson_plan", "word")
WORD_TEMPLATE_MANIFEST = os.path.join(WORD_TEMPLATE_ROOT, "manifest.json")
WORD_DEFAULT_PROFILE = "default"
WORD_DEFAULT_FONT = "Microsoft YaHei"
WORD_TITLE_COLOR = RGBColor(23, 50, 71)
WORD_HEADING_COLOR = RGBColor(55, 100, 153)


def build_output_path(upload_root: str, course_id: int, topic: str) -> str:
    materials_dir = os.path.join(upload_root, "materials", str(course_id))
    os.makedirs(materials_dir, exist_ok=True)
    base_name = secure_filename(topic) or DOCX_EXPORT_FILENAME_PREFIX
    filename = f"{base_name}_{time.strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.docx"
    return os.path.join(materials_dir, filename)


def render_docx(
    spec: Dict[str, Any],
    output_path: str,
    core_spec: Optional[Dict[str, Any]] = None,
    template_profile: Optional[str] = None,
) -> Dict[str, Any]:
    template_meta = resolve_word_template(template_profile)
    template_path = template_meta.get("template_path")
    if template_path:
        render_stats = _render_docx_from_template(
            template_path=template_path,
            spec=spec,
            output_path=output_path,
            core_spec=core_spec,
            template_profile=template_meta.get("profile"),
        )
        render_stats["template_used"] = True
        return render_stats

    document = Document()
    _configure_styles(document)

    requirement = spec.get("requirement_summary") if isinstance(spec.get("requirement_summary"), dict) else {}
    topic = str(
        requirement.get("topic")
        or requirement.get("chapter_title")
        or requirement.get("grade_subject")
        or "教案"
    ).strip()
    summary_lines = _build_summary_lines(requirement)

    title = document.add_heading(topic or "教案", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if summary_lines:
        intro = document.add_paragraph()
        intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
        intro.add_run(" | ".join(summary_lines))

    sections = spec.get("docx_outline") if isinstance(spec.get("docx_outline"), list) else []
    for item in sections:
        if not isinstance(item, dict):
            continue
        title_text = str(item.get("section_title") or "").strip()
        if not title_text:
            continue

        document.add_heading(title_text, level=1)

        section_goal = str(item.get("section_goal") or "").strip()
        if section_goal:
            goal_paragraph = document.add_paragraph()
            goal_label = goal_paragraph.add_run("章节目标：")
            goal_label.bold = True
            goal_paragraph.add_run(section_goal)

        bullets = [
            str(bullet).strip()
            for bullet in item.get("bullets", [])
            if str(bullet).strip()
        ]
        for bullet in bullets:
            document.add_paragraph(bullet, style="List Bullet")

        source_refs = [
            str(ref).strip()
            for ref in item.get("source_refs", [])
            if str(ref).strip()
        ]
        if source_refs:
            ref_paragraph = document.add_paragraph()
            ref_label = ref_paragraph.add_run("参考资料：")
            ref_label.bold = True
            ref_paragraph.add_run("；".join(source_refs))

    document.save(output_path)
    return {
        "section_count": len([item for item in sections if isinstance(item, dict) and str(item.get("section_title") or "").strip()]),
        "template_profile": template_meta.get("profile"),
        "template_used": False,
    }


def resolve_word_template(template_profile: Optional[str] = None) -> Dict[str, Any]:
    profile = str(template_profile or "").strip() or WORD_DEFAULT_PROFILE
    manifest = _load_template_manifest(WORD_TEMPLATE_MANIFEST)
    profiles = manifest.get("profiles") if isinstance(manifest.get("profiles"), dict) else {}
    profile_meta = profiles.get(profile) if isinstance(profiles.get(profile), dict) else {}
    if not profile_meta and profile != WORD_DEFAULT_PROFILE:
        profile_meta = profiles.get(WORD_DEFAULT_PROFILE) if isinstance(profiles.get(WORD_DEFAULT_PROFILE), dict) else {}
        profile = WORD_DEFAULT_PROFILE
    template_name = str(profile_meta.get("template") or f"{profile}.docx").strip()
    template_path = os.path.join(WORD_TEMPLATE_ROOT, template_name)
    if not os.path.exists(template_path):
        template_path = None
    return {"profile": profile, "template_path": template_path, "profile_meta": profile_meta}


def persist_generated_material(
    course_id: int,
    course_name: str,
    topic: str,
    output_path: str,
    section_count: int,
) -> Material:
    file_hash = _calculate_file_hash(output_path)
    existing = Material.query.filter_by(course_id=course_id, file_hash=file_hash).first()
    if existing and existing.file_path:
        existing_abs_path = _material_upload_to_abs_path(existing.file_path)
        if existing_abs_path and os.path.exists(existing_abs_path):
            if os.path.abspath(existing_abs_path) != os.path.abspath(output_path) and os.path.exists(output_path):
                os.remove(output_path)
            return existing

    filename = os.path.basename(output_path)
    relative_upload_path = f"/uploads/materials/{course_id}/{filename}"
    content = f"Generated DOCX from lesson plan spec | course={course_name} | topic={topic} | sections={section_count}"

    if existing:
        existing.title = filename
        existing.material_type = "Word"
        existing.file_path = relative_upload_path
        existing.content = content
        db.session.commit()
        return existing

    material = Material(
        title=filename,
        material_type="Word",
        file_path=relative_upload_path,
        file_hash=file_hash,
        content=content,
        course_id=course_id,
    )
    db.session.add(material)
    db.session.commit()
    return material


def _configure_styles(document: Document) -> None:
    normal_style = _find_style(document, "Normal", aliases=["正文"])
    if normal_style is not None:
        normal_style.font.name = WORD_DEFAULT_FONT
        normal_style.font.size = Pt(11)
        normal_style.font.color.rgb = WORD_TITLE_COLOR
        _set_style_font_family(normal_style, WORD_DEFAULT_FONT)

    for style_key, size, aliases in [
        ("Heading1", 16, ["Heading 1", "标题 1"]),
        ("Heading2", 13, ["Heading 2", "标题 2"]),
        ("Title", 20, ["Title", "标题"]),
    ]:
        style = _find_style(document, style_key, aliases=aliases)
        if style is None:
            continue
        style.font.name = WORD_DEFAULT_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = WORD_TITLE_COLOR if style_key == "Title" else WORD_HEADING_COLOR
        _set_style_font_family(style, WORD_DEFAULT_FONT)


def _load_template_manifest(path: str) -> Dict[str, Any]:
    if not os.path.exists(path):
        return {}
    try:
        with open(path, "r", encoding="utf-8") as file_obj:
            parsed = json.load(file_obj)
        return parsed if isinstance(parsed, dict) else {}
    except Exception:
        return {}


def _render_docx_from_template(
    *,
    template_path: str,
    spec: Dict[str, Any],
    output_path: str,
    core_spec: Optional[Dict[str, Any]],
    template_profile: Optional[str],
) -> Dict[str, Any]:
    document = Document(template_path)
    _configure_styles(document)
    requirement = spec.get("requirement_summary") if isinstance(spec.get("requirement_summary"), dict) else {}
    core_spec = core_spec if isinstance(core_spec, dict) else {}
    teaching_objectives = core_spec.get("teaching_objectives") if isinstance(core_spec.get("teaching_objectives"), dict) else {}
    student_profile = core_spec.get("student_profile") if isinstance(core_spec.get("student_profile"), dict) else {}
    teaching_flow = core_spec.get("teaching_flow") if isinstance(core_spec.get("teaching_flow"), list) else []
    assessment_plan = core_spec.get("assessment_plan") if isinstance(core_spec.get("assessment_plan"), dict) else {}
    knowledge_structure = core_spec.get("knowledge_structure") if isinstance(core_spec.get("knowledge_structure"), dict) else {}
    if not teaching_flow:
        teaching_flow = _build_teaching_flow_from_docx_outline(spec)

    topic = str(
        requirement.get("topic")
        or requirement.get("chapter_title")
        or requirement.get("grade_subject")
        or "教案"
    ).strip()

    objectives = _normalize_sentence_list(
        teaching_objectives.get("goals")
    ) or _normalize_sentence_list(requirement.get("teaching_goals"))
    process_goals = _normalize_sentence_list(
        teaching_objectives.get("deliverables")
    ) or [
        "通过案例推演、板书梳理和课堂互动帮助学生形成清晰的方法路径，并能在新情境中迁移使用。",
        "通过当堂任务练习和同伴交流，逐步形成分析问题、组织表达和反思修正的学习方法。"
    ]
    value_goals = _normalize_sentence_list(knowledge_structure.get("examples_or_cases")) or ["建立规范表达与反思意识，并能够在课堂任务中主动展示自己的理解过程。"]
    key_points = _normalize_sentence_list(teaching_objectives.get("key_points")) or _normalize_sentence_list(requirement.get("key_points"))
    difficult_points = _normalize_sentence_list(teaching_objectives.get("difficult_points")) or _normalize_sentence_list(requirement.get("difficult_points"))
    board_design = _build_board_design(core_spec, requirement)
    homework_lines = _build_homework_lines(core_spec)
    reflection_lines = _build_reflection_lines(core_spec)
    checks = _normalize_sentence_list(assessment_plan.get("questions")) or _normalize_sentence_list(assessment_plan.get("in_class_checks"))
    feedback_text = _build_feedback_text(assessment_plan, topic)
    student_summary = _build_student_summary(student_profile, requirement, topic)

    paragraphs = document.paragraphs
    heading_indices = {
        "一、课题": _find_paragraph_index(paragraphs, "一、课题"),
        "（一）知识与技能目标": _find_paragraph_index(paragraphs, "（一）知识与技能目标"),
        "（二）过程与方法目标": _find_paragraph_index(paragraphs, "（二）过程与方法目标"),
        "（三）情感态度与价值观目标": _find_paragraph_index(paragraphs, "（三）情感态度与价值观目标"),
        "三、学情分析": _find_paragraph_index(paragraphs, "三、学情分析"),
        "（一）教学重点": _find_paragraph_index(paragraphs, "（一）教学重点"),
        "（二）教学难点": _find_paragraph_index(paragraphs, "（二）教学难点"),
        "（一）检测题目": _find_paragraph_index(paragraphs, "（一）检测题目"),
        "（二）检测反馈": _find_paragraph_index(paragraphs, "（二）检测反馈"),
        "七、板书设计": _find_paragraph_index(paragraphs, "七、板书设计"),
        "（一）课后作业": _find_paragraph_index(paragraphs, "（一）课后作业"),
        "（二）教学反思": _find_paragraph_index(paragraphs, "（二）教学反思"),
    }

    metadata_map = {
        "学校抬头：": str(requirement.get("school_name") or "__________________________").strip(),
        "授课教师：": str(requirement.get("teacher_name") or "________________________").strip(),
        "授课年级：": str(student_profile.get("grade") or "________________________").strip(),
        "学科：": str(requirement.get("grade_subject") or "__________________________").strip(),
        "授课时长：": str(requirement.get("duration") or "________________________").strip(),
        "授课日期：": str(requirement.get("teach_date") or "______年____月____日").strip(),
        "授课班级：": str(requirement.get("class_name") or "________________________").strip(),
    }
    _fill_metadata_lines(paragraphs, metadata_map)

    _fill_placeholder_paragraphs(paragraphs, heading_indices["一、课题"], [_build_topic_section_text(requirement, objectives)])
    _fill_placeholder_paragraphs(paragraphs, heading_indices["（一）知识与技能目标"], objectives, bullet=True)
    _fill_placeholder_paragraphs(paragraphs, heading_indices["（二）过程与方法目标"], process_goals, bullet=True)
    _fill_placeholder_paragraphs(paragraphs, heading_indices["（三）情感态度与价值观目标"], value_goals, bullet=True)
    _fill_placeholder_paragraphs(paragraphs, heading_indices["三、学情分析"], [student_summary])
    _fill_placeholder_paragraphs(paragraphs, heading_indices["（一）教学重点"], key_points, bullet=True)
    _fill_placeholder_paragraphs(paragraphs, heading_indices["（二）教学难点"], difficult_points, bullet=True)
    _fill_placeholder_paragraphs(paragraphs, heading_indices["（一）检测题目"], checks[:2], bullet=True)
    _expand_bullet_section(
        paragraphs,
        heading_indices["（一）检测题目"],
        heading_indices["（二）检测反馈"],
        checks,
        bullet=True,
    )
    _fill_placeholder_paragraphs(paragraphs, heading_indices["（二）检测反馈"], [feedback_text])
    _fill_placeholder_paragraphs(paragraphs, heading_indices["七、板书设计"], board_design)
    _fill_homework_lines(paragraphs, heading_indices["（一）课后作业"], homework_lines)
    _fill_placeholder_paragraphs(paragraphs, heading_indices["（二）教学反思"], reflection_lines)

    if document.tables and not _render_teaching_process_as_paragraphs(document, document.tables[0], teaching_flow):
        _fill_teaching_process_table(document.tables[0], teaching_flow)

    _strip_template_scaffolding(document)
    _rebuild_title_block(document, requirement)
    _normalize_template_paragraph_styles(document)
    _normalize_metadata_paragraphs(document)

    document.save(output_path)
    return {
        "section_count": len([item for item in spec.get("docx_outline", []) if isinstance(item, dict) and str(item.get("section_title") or "").strip()]),
        "template_profile": template_profile or WORD_DEFAULT_PROFILE,
    }


def _find_paragraph_index(paragraphs: List[Any], marker: str) -> int:
    for index, paragraph in enumerate(paragraphs):
        if marker in str(paragraph.text or ""):
            return index
    return -1


def _normalize_list(raw_value: Any) -> List[str]:
    if isinstance(raw_value, list):
        return [str(item).strip() for item in raw_value if str(item).strip()]
    if isinstance(raw_value, str):
        text = raw_value.strip()
        return [text] if text else []
    return []


def _normalize_sentence_list(raw_value: Any) -> List[str]:
    items = _normalize_list(raw_value)
    results: List[str] = []
    seen = set()
    for item in items:
        text = str(item or "").strip()
        if not text:
            continue
        if len(text) < 36:
            text = f"{text}，教师需在课堂中结合示例讲解、提问互动和即时反馈，帮助学生真正理解并能够迁移应用。"
        if text not in seen:
            seen.add(text)
            results.append(text)
    return results


def _set_paragraph_text(paragraph: Any, text: str) -> None:
    _replace_paragraph_runs(paragraph, str(text or ""))
    _clear_paragraph_numbering(paragraph)
    _apply_run_font_overrides(paragraph)


def _fill_metadata_lines(paragraphs: List[Any], metadata_map: Dict[str, str]) -> None:
    for paragraph in paragraphs:
        raw_text = str(paragraph.text or "")
        if not raw_text.strip():
            continue
        updated_text = raw_text
        for label, value in metadata_map.items():
            if label in updated_text:
                updated_text = _replace_label_placeholder(updated_text, label, value)
        if updated_text != raw_text:
            _set_paragraph_text(paragraph, updated_text)


def _replace_label_placeholder(text: str, label: str, value: str) -> str:
    pattern = re.escape(label) + r"[^\s（）\(]*"
    match = re.search(pattern, text)
    if not match:
        return text
    return text[:match.start()] + f"{label}{value}" + text[match.end():]


def _strip_template_scaffolding(document: Document) -> None:
    paragraphs = document.paragraphs
    metadata_index = _find_paragraph_index(paragraphs, "学校抬头：")
    if metadata_index > 0:
        for index in range(metadata_index - 1, -1, -1):
            _remove_paragraph(paragraphs[index])

    for paragraph in list(document.paragraphs):
        raw_text = str(paragraph.text or "")
        text = raw_text.strip()
        if not text:
            continue

        if text.startswith("（表格样式：") or text.startswith("样式说明："):
            _remove_paragraph(paragraph)
            continue

        cleaned = re.sub(r"（样式：[^）]*）", "", raw_text)
        cleaned = re.sub(r"（此处由软件填充[^）]*）", "", cleaned)
        cleaned = re.sub(r"（软件填充[^）]*）", "", cleaned)
        cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
        if cleaned != raw_text.strip():
            _set_paragraph_text(paragraph, cleaned)


def _rebuild_title_block(document: Document, requirement: Dict[str, Any]) -> None:
    paragraphs = document.paragraphs
    metadata_index = _find_paragraph_index(paragraphs, "授课教师：")
    if metadata_index < 0:
        return

    topic = str(
        requirement.get("topic")
        or requirement.get("chapter_title")
        or requirement.get("grade_subject")
        or "教案"
    ).strip()
    summary_lines = _build_summary_lines(requirement)

    anchor = document.paragraphs[_find_paragraph_index(document.paragraphs, "授课教师：")]
    summary_paragraph = None
    if summary_lines:
        summary_paragraph = anchor.insert_paragraph_before(" | ".join(summary_lines))
        summary_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        summary_paragraph.paragraph_format.space_after = Pt(12)
        summary_paragraph.paragraph_format.line_spacing = 1.2
        _apply_run_font_overrides(summary_paragraph, size=11, bold=False, color=WORD_HEADING_COLOR)

    title_paragraph = (summary_paragraph or anchor).insert_paragraph_before(topic)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(8)
    title_paragraph.paragraph_format.line_spacing = 1.2
    _apply_run_font_overrides(title_paragraph, size=22, bold=True, color=WORD_TITLE_COLOR)


def _build_topic_section_text(requirement: Dict[str, Any], objectives: List[str]) -> str:
    topic = str(
        requirement.get("topic")
        or requirement.get("chapter_title")
        or requirement.get("grade_subject")
        or "本课主题"
    ).strip()
    goal_text = "；".join([str(item).strip() for item in objectives[:3] if str(item).strip()])
    if goal_text:
        return f"本课题围绕“{topic}”展开，教师将通过问题导入、概念讲解、案例分析、课堂互动与当堂练习，引导学生在真实任务中逐步理解核心知识、形成分析方法，并落实以下目标：{goal_text}"
    return f"本课题围绕“{topic}”展开，教师将结合课堂讲解、案例引导、结构化板书与当堂练习，帮助学生逐步建立完整的知识理解、方法迁移与实际应用能力。"


def _build_student_summary(student_profile: Dict[str, Any], requirement: Dict[str, Any], topic: str) -> str:
    foundation = str(student_profile.get("foundation") or "").strip()
    preference = str(student_profile.get("learning_preference") or "").strip()
    misconceptions = _normalize_sentence_list(student_profile.get("common_misconceptions"))
    support = _normalize_sentence_list(student_profile.get("support_strategies"))
    grade = str(student_profile.get("grade") or requirement.get("grade_subject") or "当前学生").strip()

    parts: List[str] = []
    if foundation:
        parts.append(f"{grade}学生当前基础情况为：{foundation}")
    if preference:
        parts.append(f"在学习方式上更适合采用{preference}的组织形式，通过讲练结合、分层提问和即时反馈提升参与度")
    if misconceptions:
        parts.append(f"围绕“{topic}”的学习中，学生常见的理解偏差包括：{'；'.join(misconceptions[:2])}")
    if support:
        parts.append(f"课堂支持策略建议为：{'；'.join(support[:2])}")

    if parts:
        return "；".join(parts)
    return f"学生对“{topic}”具备一定生活经验或前置知识，但容易停留在表层记忆。教师需要通过案例类比、结构化板书、分层提问和即时练习，帮助学生把零散认识转化为可表达、可迁移的系统理解。"


def _build_feedback_text(assessment_plan: Dict[str, Any], topic: str) -> str:
    checks = _normalize_sentence_list(assessment_plan.get("in_class_checks"))
    questions = _normalize_sentence_list(assessment_plan.get("questions"))
    merged = checks + [item for item in questions if item not in checks]
    if merged:
        return f"课堂检测后，教师应结合以下观察点进行针对性反馈与错因追踪：{'；'.join(merged[:3])}"
    return f"课堂检测后，教师需要围绕“{topic}”及时归纳学生易错点，结合典型问题进行二次讲解，并通过追问、板演或当堂订正帮助学生把关键概念真正落实到位。"


def _fill_placeholder_paragraphs(paragraphs: List[Any], heading_index: int, values: List[str], bullet: bool = False) -> None:
    if heading_index < 0:
        return
    placeholder_indexes: List[int] = []
    for index in range(heading_index + 1, len(paragraphs)):
        text = str(paragraphs[index].text or "").strip()
        if not text:
            continue
        if re.match(r"^[一二三四五六七八九十]+、", text) or re.match(r"^（[一二三四五六七八九十]+）", text):
            break
        if "________________" in text:
            placeholder_indexes.append(index)
    if not placeholder_indexes:
        return
    if not values:
        values = ["待补充"]
    for slot, paragraph_index in enumerate(placeholder_indexes):
        paragraph = paragraphs[paragraph_index]
        text = values[slot] if slot < len(values) else ""
        if not text:
            _set_paragraph_text(paragraph, "")
            continue
        _set_paragraph_text(paragraph, text if not bullet else f"• {text}")
        if bullet:
            _configure_manual_bullet_paragraph(paragraph)


def _fill_homework_lines(paragraphs: List[Any], heading_index: int, homework_lines: Dict[str, str]) -> None:
    if heading_index < 0:
        return
    for index in range(heading_index + 1, len(paragraphs)):
        text = str(paragraphs[index].text or "").strip()
        if not text:
            continue
        if re.match(r"^（二）教学反思", text):
            break
        if text.startswith("基础作业："):
            _set_paragraph_text(paragraphs[index], f"基础作业：{homework_lines['basic']}")
        elif text.startswith("拓展作业："):
            _set_paragraph_text(paragraphs[index], f"拓展作业：{homework_lines['extension']}")


def _expand_bullet_section(
    paragraphs: List[Any],
    heading_index: int,
    end_heading_index: int,
    values: List[str],
    *,
    bullet: bool = True,
) -> None:
    if heading_index < 0 or end_heading_index < 0:
        return
    if len(values) <= 2:
        return
    insert_before = paragraphs[end_heading_index]
    for text in values[2:][::-1]:
        if not str(text or "").strip():
            continue
        paragraph = _insert_paragraph_before(insert_before, f"• {text}" if bullet else str(text))
        if bullet:
            _configure_manual_bullet_paragraph(paragraph)
        else:
            _apply_run_font_overrides(paragraph)


def _build_homework_lines(core_spec: Dict[str, Any]) -> Dict[str, str]:
    assessment_plan = core_spec.get("assessment_plan") if isinstance(core_spec.get("assessment_plan"), dict) else {}
    homework = _normalize_sentence_list(assessment_plan.get("homework"))
    extension = _normalize_sentence_list(assessment_plan.get("extension_tasks"))
    return {
        "basic": "；".join(homework[:3]) or "完成课堂知识点整理与基础练习，并结合课堂板书梳理本节课的核心概念、关键术语和方法步骤，确保能够独立复述主要内容。",
        "extension": "；".join(extension[:3]) or "结合真实情境完成迁移拓展任务，尝试把课堂所学用于解释现象、分析案例或解决新的问题，并记录自己的思考过程。",
    }


def _build_reflection_lines(core_spec: Dict[str, Any]) -> List[str]:
    grounding = core_spec.get("source_grounding") if isinstance(core_spec.get("source_grounding"), list) else []
    claims = [str(item.get("claim") or "").strip() for item in grounding if isinstance(item, dict) and str(item.get("claim") or "").strip()]
    return _normalize_sentence_list(claims[:3]) or [
        "复盘课堂互动效果与学生掌握情况，特别关注学生在哪些环节能够主动表达、在哪些环节仍需要教师支架支持。",
        "根据作业反馈和课堂检测结果调整后续教学节奏、例题难度和练习层次，提高教学针对性。",
        "结合本节课的时间分配、问题设计和活动组织方式，评估哪些教学支架真正帮助了学生理解，哪些环节仍需要进一步优化。"
    ]


def _build_board_design(core_spec: Dict[str, Any], requirement: Dict[str, Any]) -> List[str]:
    knowledge_structure = core_spec.get("knowledge_structure") if isinstance(core_spec.get("knowledge_structure"), dict) else {}
    points = _normalize_list(knowledge_structure.get("knowledge_points")) or _normalize_list(requirement.get("knowledge_points"))
    lines = [
        f"{index + 1}. {item}，板书时建议同步标注关键词、概念关系和典型示例，帮助学生形成清晰的知识结构。"
        for index, item in enumerate(points[:3])
    ]
    return lines or [
        "1. 主题导入，明确本节课问题情境、学习目标和核心任务。",
        "2. 核心知识，按主概念、关键方法和典型案例的顺序展开板书。",
        "3. 总结提升，用结构化方式回顾重点难点并提示课后迁移方向。"
    ]


def _fill_teaching_process_table(table: Any, teaching_flow: List[Dict[str, Any]]) -> None:
    flows = [item for item in teaching_flow if isinstance(item, dict)]
    table.autofit = False
    _set_table_column_widths(table, [1.15, 0.75, 1.95, 1.9])
    while len(table.rows) - 1 < len(flows):
        table.add_row()
    for row_index, flow in enumerate(flows, start=1):
        cells = table.rows[row_index].cells
        activities = _compact_flow_text(flow.get("activities"), fallback=flow.get("goal"), max_chars=24)
        teacher_actions = _compact_flow_text(flow.get("teacher_actions"), fallback=activities, max_chars=18)
        student_actions = _compact_flow_text(flow.get("student_actions"), fallback="", max_chars=18)
        assessments = _compact_flow_text(flow.get("assessment"), fallback="", max_chars=18)
        design_intent = _compact_flow_text(flow.get("goal"), fallback="", max_chars=24)
        execution_text = "；".join(
            part for part in [
                f"师：{teacher_actions}" if teacher_actions else "",
                f"生：{student_actions}" if student_actions else "",
                activities if activities and activities not in {teacher_actions, student_actions} else "",
            ]
            if part
        )
        intent_text = "；".join(
            part for part in [
                f"目标：{design_intent}" if design_intent else "",
                f"评价：{assessments}" if assessments else "",
            ]
            if part
        )
        cells[0].text = str(flow.get("title") or f"教学环节{row_index}").strip()
        cells[1].text = _estimate_flow_duration(row_index, len(flows))
        cells[2].text = execution_text or activities or "教师组织活动，学生参与完成任务。"
        cells[3].text = intent_text or "围绕本环节目标组织活动，并根据学生表现及时调整讲解与练习安排。"
        _style_table_text_cell(cells[0], size=10)
        _style_table_text_cell(cells[1], size=10)
        _style_table_text_cell(cells[2], size=10)
        _style_table_text_cell(cells[3], size=10)
    for row_index in range(len(flows) + 1, len(table.rows)):
        for cell in table.rows[row_index].cells:
            cell.text = ""


def _render_teaching_process_as_paragraphs(document: Document, table: Any, teaching_flow: List[Dict[str, Any]]) -> bool:
    flows = [item for item in teaching_flow if isinstance(item, dict)]
    if not flows:
        return False

    next_heading = None
    for paragraph in document.paragraphs:
        if "六、课堂检测" in str(paragraph.text or ""):
            next_heading = paragraph
            break
    if next_heading is None:
        return False

    for flow_index, flow in enumerate(flows, start=1):
        stage_title = str(flow.get("title") or "教学环节").strip()
        duration = str(flow.get("duration") or "").strip() or _estimate_flow_duration(flow_index, len(flows))
        activities = "；".join(_normalize_list(flow.get("activities"))) or str(flow.get("goal") or "").strip() or "组织课堂活动"
        teacher_actions = "；".join(_normalize_list(flow.get("teacher_actions")))
        student_actions = "；".join(_normalize_list(flow.get("student_actions")))
        assessments = "；".join(_normalize_list(flow.get("assessment")))
        goal = str(flow.get("goal") or "").strip()

        title_paragraph = _insert_paragraph_before(next_heading, f"{stage_title}（{duration}）")
        title_paragraph.style = _find_style(document, "Heading2", aliases=["标题 2"]) or title_paragraph.style
        _apply_run_font_overrides(title_paragraph, size=12, bold=True, color=WORD_HEADING_COLOR)

        if activities:
            activity_paragraph = _insert_paragraph_before(next_heading, f"教学活动：{activities}")
            _apply_run_font_overrides(activity_paragraph, size=10, bold=False)
        if teacher_actions:
            teacher_paragraph = _insert_paragraph_before(next_heading, f"教师活动：{teacher_actions}")
            _apply_run_font_overrides(teacher_paragraph, size=10, bold=False)
        if student_actions:
            student_paragraph = _insert_paragraph_before(next_heading, f"学生活动：{student_actions}")
            _apply_run_font_overrides(student_paragraph, size=10, bold=False)
        if goal or assessments:
            intent_text = "；".join(
                part for part in [
                    f"设计意图：{goal}" if goal else "",
                    f"评价提示：{assessments}" if assessments else "",
                ]
                if part
            )
            intent_paragraph = _insert_paragraph_before(next_heading, intent_text)
            _apply_run_font_overrides(intent_paragraph, size=10, bold=False)

        spacer = _insert_paragraph_before(next_heading, "")
        spacer.paragraph_format.space_after = Pt(4)

    _remove_table(table)
    for paragraph in list(document.paragraphs):
        text = str(paragraph.text or "").strip()
        if "表格样式：统一设置为“三线表”" in text or text.startswith("（表格样式：统一设置为"):
            _remove_paragraph(paragraph)
    return True


def _insert_paragraph_before(paragraph: Paragraph, text: str, style: Optional[Any] = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def _remove_table(table: Any) -> None:
    element = table._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _chunk_list(items: List[str], chunk_size: int) -> List[List[str]]:
    if chunk_size <= 0:
        chunk_size = 1
    return [items[index:index + chunk_size] for index in range(0, len(items), chunk_size)]


def _guess_flow_title_from_text(text: str, index: int) -> str:
    normalized = str(text or "").strip()
    title_rules = [
        ("导入", "导入新课"),
        ("情境", "情境导入"),
        ("探究", "合作探究"),
        ("讲授", "新知讲授"),
        ("新知", "新知讲授"),
        ("练习", "巩固练习"),
        ("活动", "课堂活动"),
        ("总结", "课堂小结"),
        ("作业", "迁移拓展"),
    ]
    for keyword, title in title_rules:
        if keyword in normalized:
            return title
    fallback_titles = ["导入新课", "新知讲授", "巩固练习", "课堂小结", "迁移拓展"]
    return fallback_titles[min(index, len(fallback_titles) - 1)]


def _build_teaching_flow_from_docx_outline(spec: Dict[str, Any]) -> List[Dict[str, Any]]:
    if not isinstance(spec, dict):
        return []
    docx_outline = spec.get("docx_outline") if isinstance(spec.get("docx_outline"), list) else []
    candidate_sections = [
        item for item in docx_outline
        if isinstance(item, dict) and any(keyword in str(item.get("section_title") or "") for keyword in ["教学流程", "课堂活动"])
    ]
    if not candidate_sections:
        return []

    flows: List[Dict[str, Any]] = []
    for section in candidate_sections:
        section_goal = str(section.get("section_goal") or "").strip()
        bullets = _normalize_list(section.get("bullets"))
        if not bullets:
            continue
        chunk_size = 2 if len(bullets) > 4 else 1
        bullet_groups = _chunk_list(bullets, chunk_size)
        for index, group in enumerate(bullet_groups):
            joined = "；".join(group)
            flows.append({
                "title": _guess_flow_title_from_text(joined or section_goal, len(flows)),
                "goal": section_goal or (group[0] if group else "组织课堂活动"),
                "activities": group[:2],
                "teacher_actions": [group[0]] if group else [],
                "student_actions": [group[1]] if len(group) > 1 else [],
                "assessment": [group[-1]] if group else [],
            })
            if len(flows) >= 5:
                return flows
    return flows


def _set_table_column_widths(table: Any, widths: List[float]) -> None:
    columns = list(getattr(table, "columns", []))
    for index, width in enumerate(widths):
        if index >= len(columns):
            break
        columns[index].width = Inches(width)
    for row in table.rows:
        for index, width in enumerate(widths):
            if index >= len(row.cells):
                break
            row.cells[index].width = Inches(width)


def _compact_flow_text(raw_value: Any, *, fallback: Any = "", max_chars: int = 24) -> str:
    candidates = _normalize_list(raw_value)
    if not candidates:
        candidates = _normalize_list(fallback)
    if not candidates and isinstance(fallback, str):
        candidates = [fallback.strip()] if fallback.strip() else []
    return "；".join(candidates) if candidates else ""


def _style_table_text_cell(cell: Any, *, size: int = 10) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        _apply_run_font_overrides(paragraph, size=size, bold=False)


def _set_style_font_family(style: Any, font_name: str) -> None:
    rpr = getattr(style._element, "rPr", None)
    if rpr is None:
        rpr = style._element.get_or_add_rPr()
    r_fonts = rpr.rFonts
    if r_fonts is None:
        r_fonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(attr), font_name)


def _find_style(document: Document, style_id: str, aliases: Optional[List[str]] = None) -> Optional[Any]:
    aliases = aliases or []
    for style in document.styles:
        if str(getattr(style, "style_id", "") or "") == style_id:
            return style
        style_name = str(getattr(style, "name", "") or "")
        if style_name in aliases:
            return style
    return None


def _replace_paragraph_runs(paragraph: Any, text: str) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    if text:
        paragraph.add_run(text)


def _remove_paragraph(paragraph: Any) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def _clear_paragraph_numbering(paragraph: Any) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    num_pr = p_pr.numPr
    if num_pr is not None:
        p_pr.remove(num_pr)


def _configure_manual_bullet_paragraph(paragraph: Any) -> None:
    _clear_paragraph_numbering(paragraph)
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-10)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.45


def _apply_font_to_run(run: Any, *, size: Optional[int] = None, bold: Optional[bool] = None, color: Optional[RGBColor] = None) -> None:
    run.font.name = WORD_DEFAULT_FONT
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(attr), WORD_DEFAULT_FONT)


def _apply_run_font_overrides(
    paragraph: Any,
    *,
    size: Optional[int] = None,
    bold: Optional[bool] = None,
    color: Optional[RGBColor] = None,
) -> None:
    for run in paragraph.runs:
        _apply_font_to_run(run, size=size, bold=bold, color=color)


def _normalize_template_paragraph_styles(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = str(paragraph.text or "").strip()
        style_name = str(getattr(paragraph.style, "name", "") or "")
        style_id = str(getattr(paragraph.style, "style_id", "") or "")
        _clear_paragraph_numbering(paragraph)
        paragraph.paragraph_format.line_spacing = 1.45
        if re.match(r"^（[一二三四五六七八九十]+）", text):
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
        elif re.match(r"^[一二三四五六七八九十]+、", text):
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
        else:
            paragraph.paragraph_format.space_after = Pt(6)
        if style_id == "Title" or style_name == "标题":
            _apply_run_font_overrides(paragraph, size=20, bold=True, color=WORD_TITLE_COLOR)
        elif style_id == "Heading1" or style_name == "标题 1":
            _apply_run_font_overrides(paragraph, size=16, bold=True, color=WORD_HEADING_COLOR)
        elif style_id == "Heading2" or style_name == "标题 2":
            _apply_run_font_overrides(paragraph, size=13, bold=True, color=WORD_HEADING_COLOR)
        elif re.match(r"^[一二三四五六七八九十]+、", text):
            _apply_run_font_overrides(paragraph, size=16, bold=True, color=WORD_HEADING_COLOR)
        elif re.match(r"^（[一二三四五六七八九十]+）", text):
            _apply_run_font_overrides(paragraph, size=13, bold=True, color=WORD_HEADING_COLOR)


def _normalize_metadata_paragraphs(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = str(paragraph.text or "").strip()
        if not text:
            continue
        if text.startswith("学校抬头："):
            _remove_paragraph(paragraph)
            continue
        if any(text.startswith(prefix) for prefix in ("授课教师：", "学科：", "授课日期：")):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            _apply_run_font_overrides(paragraph, size=11, bold=False, color=WORD_TITLE_COLOR)


def _estimate_flow_duration(index: int, total: int) -> str:
    if total <= 0:
        return "10分钟"
    if index == 1:
        return "8分钟"
    if index == total:
        return "7分钟"
    return "10分钟"


def _build_summary_lines(requirement: Dict[str, Any]) -> List[str]:
    profile = requirement.get("student_profile") if isinstance(requirement.get("student_profile"), dict) else {}
    return [
        item
        for item in [
            str(requirement.get("grade_subject") or "").strip(),
            str(requirement.get("chapter_title") or "").strip(),
            str(requirement.get("duration") or "").strip(),
            str(profile.get("foundation") or "").strip(),
        ]
        if item
    ]


def _calculate_file_hash(file_path: str) -> Optional[str]:
    try:
        digest = hashlib.sha256()
        with open(file_path, "rb") as file_obj:
            for chunk in iter(lambda: file_obj.read(4096), b""):
                digest.update(chunk)
        return digest.hexdigest()
    except Exception:
        return None


def _material_upload_to_abs_path(file_path: str) -> Optional[str]:
    if not isinstance(file_path, str) or not file_path:
        return None
    project_backend_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(project_backend_dir, file_path.lstrip("/"))
